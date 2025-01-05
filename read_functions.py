### read_functions.py
## Functions to be imported to main script
import pandas as pd
import pypdfium2 as pdfium
import numpy as np
import docx
import sys
###########

##
## Convert Word (docx) file to Pandas data format
def read_docx(filen):
	# Load the Word document
	doc = docx.Document(filen)
	
	# Access the table
	table = doc.tables[0]
	
	# Create an empty list to store lines in
	data = []
	
	# Iterate through rows in the table
	for row in table.rows:
		row_data = []
		for cell in row.cells:
			row_data.append(cell.text)
		data.append(row_data)
	
	# Convert list to Pandas data frame
	df = pd.DataFrame(data)
	
	# Set first row as header and restructure data frame
	# // Maybe not necessary, or at least not yet?
	# ~ df.columns = df.iloc[0]
	# ~ df = df[1:]
	# ~ df.reset_index(drop=True, inplace=True)
	return df

##
## Function to read PDF files 
def read_pdf(fname):
	pdf = pdfium.PdfDocument(fname)
	n_pages =len(pdf)

	pdf_text_raw = np.array([])	# Requires array for data structure later
	for i in range(n_pages):
		page = pdf[i]	# Has to read each page individually, but handles multiple pages now
		pagetext = (page.get_textpage()).get_text_range(force_this=True)	# No errors
		pagerows = pagetext.split('\n')
		pdf_text_raw = np.append(pdf_text_raw, pagerows)	# Clean artefacts later

	# Cleaning some artefacts
	pdf_text = []
	for lines in pdf_text_raw:
		pdf_text.append(lines.replace("\r", ""))
	
	df = pd.DataFrame(pdf_text)		# Panda Data structure becuase we worked with before
	return df

##
## Main function to search words
def search_words(df):
	Q = {}    # Make a dictionary
	for index, row in df.iterrows():
		# Step through rows in table
		for column in df.columns:
			value = str(row[column])
			# Save ID
			if "ID" in value:
				current_ID = value
				answers = []	# Save all four answers in one list
				continue		# skip to next row
			
			# Find and save words in both questions and answers
			if "Question:" in value:
				answers.append(value[10:].casefold())	# Ensure answer is all lower case
			if ")" in value:
				answers.append(value[3:].casefold())	
			if len(answers) == 5:
				Q.update({current_ID : answers})
			if "4)" in value:			
				continue				# skip to next row		
	
	return Q




